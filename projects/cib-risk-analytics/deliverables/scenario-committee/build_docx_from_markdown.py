from __future__ import annotations

import re
from pathlib import Path

import markdown
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_COLOR_INDEX, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BASE = Path(__file__).resolve().parent
FILES = [
    'AI_Scenario_Committee_Report_2026-03-06.md',
    'AI_Scenario_Committee_One_Pager_2026-03-06.md',
    'AI_Scenario_Committee_Speaking_Script_2026-03-06.md',
]

ACCENT = RGBColor(31, 78, 121)
ACCENT_LIGHT = 'DCE6F1'
ACCENT_MID = 'B8CCE4'
GRAY_LIGHT = 'F3F4F6'
GRAY_BORDER = 'C7CED6'
TEXT_DARK = RGBColor(36, 36, 36)
TEXT_MUTED = RGBColor(90, 90, 90)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_border(cell, color: str = GRAY_BORDER, size: str = '6') -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn('w:tcBorders'))
    if tc_borders is None:
        tc_borders = OxmlElement('w:tcBorders')
        tc_pr.append(tc_borders)
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = qn(f'w:{edge}')
        element = tc_borders.find(tag)
        if element is None:
            element = OxmlElement(f'w:{edge}')
            tc_borders.append(element)
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), size)
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), color)


def set_paragraph_border(paragraph, color: str = ACCENT_MID, size: str = '6') -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    pbdr = p_pr.find(qn('w:pBdr'))
    if pbdr is None:
        pbdr = OxmlElement('w:pBdr')
        p_pr.append(pbdr)
    bottom = pbdr.find(qn('w:bottom'))
    if bottom is None:
        bottom = OxmlElement('w:bottom')
        pbdr.append(bottom)
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = 'PAGE'
    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    r_pr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '1F4E79')
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    r_pr.append(color)
    r_pr.append(underline)
    new_run.append(r_pr)
    text_el = OxmlElement('w:t')
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def configure_document(doc: Document, title: str, kind: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    if kind == 'one_pager':
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
    else:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = kind == 'report'

    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Georgia'
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TEXT_DARK
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for style_name, size, color, space_before, space_after in [
        ('Heading 1', 15, ACCENT, 18, 8),
        ('Heading 2', 12.5, ACCENT, 14, 6),
        ('Heading 3', 11.5, ACCENT, 12, 4),
    ]:
        style = styles[style_name]
        style.font.name = 'Arial'
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(space_before)
        style.paragraph_format.space_after = Pt(space_after)
        style.paragraph_format.keep_with_next = True

    if 'Doc Title' not in styles:
        style = styles.add_style('Doc Title', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Arial'
        style.font.size = Pt(24)
        style.font.bold = True
        style.font.color.rgb = ACCENT
        style.paragraph_format.space_after = Pt(10)
    if 'Doc Subtitle' not in styles:
        style = styles.add_style('Doc Subtitle', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        style.font.color.rgb = TEXT_MUTED
        style.paragraph_format.space_after = Pt(10)
    if 'Lead Paragraph' not in styles:
        style = styles.add_style('Lead Paragraph', WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles['Normal']
        style.font.size = Pt(11)
        style.font.italic = True
        style.font.color.rgb = TEXT_MUTED
        style.paragraph_format.space_after = Pt(10)
    if 'Quote Custom' not in styles:
        style = styles.add_style('Quote Custom', WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles['Normal']
        style.font.italic = True
        style.font.color.rgb = TEXT_MUTED
        style.paragraph_format.left_indent = Inches(0.25)
        style.paragraph_format.right_indent = Inches(0.1)
        style.paragraph_format.space_after = Pt(8)
    if 'Code Block' not in styles:
        style = styles.add_style('Code Block', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Consolas'
        style.font.size = Pt(9.5)
        style.paragraph_format.left_indent = Inches(0.18)
        style.paragraph_format.right_indent = Inches(0.1)
        style.paragraph_format.space_after = Pt(8)
    if 'Small Note' not in styles:
        style = styles.add_style('Small Note', WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(8.5)
        style.font.color.rgb = TEXT_MUTED
        style.paragraph_format.space_after = Pt(4)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.style = styles['Small Note']
    header.add_run(title)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.style = styles['Small Note']
    footer.add_run('AI scenario committee pack | Page ')
    add_page_number(footer)

    props = doc.core_properties
    props.title = title
    props.subject = 'AI scenario committee materials'
    props.author = 'Codex'


def markdown_to_soup(text: str) -> BeautifulSoup:
    html = markdown.markdown(text, extensions=['tables', 'fenced_code', 'sane_lists', 'nl2br'])
    return BeautifulSoup(html, 'lxml')


def parse_metadata_paragraph(tag: Tag) -> list[tuple[str, str]]:
    if tag.name != 'p':
        return []
    chunks = re.split(r'<br\s*/?>', tag.decode_contents())
    items: list[tuple[str, str]] = []
    for chunk in chunks:
        if not chunk.strip():
            continue
        line = BeautifulSoup(chunk, 'lxml')
        strong = line.find(['strong', 'b'])
        line_text = line.get_text(' ', strip=True)
        if strong is None:
            continue
        label = strong.get_text(' ', strip=True).rstrip(':').strip()
        strong.extract()
        value = line.get_text(' ', strip=True).lstrip(':').strip()
        if label and value:
            items.append((label, value))
    return items if len(items) >= 2 else []


def append_inline(paragraph, node, bold: bool = False, italic: bool = False, code: bool = False) -> None:
    if isinstance(node, NavigableString):
        text = str(node)
        if not text:
            return
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        if code:
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.highlight_color = WD_COLOR_INDEX.GRAY_25
        return

    if not isinstance(node, Tag):
        return

    name = node.name.lower()
    if name == 'br':
        paragraph.add_run().add_break(WD_BREAK.LINE)
        return
    if name == 'a':
        add_hyperlink(paragraph, node.get_text(' ', strip=True), node.get('href', ''))
        return
    next_bold = bold or name in {'strong', 'b'}
    next_italic = italic or name in {'em', 'i'}
    next_code = code or name == 'code'
    for child in node.children:
        append_inline(paragraph, child, next_bold, next_italic, next_code)


def add_metadata_table(doc: Document, items: list[tuple[str, str]], center: bool = False) -> None:
    table = doc.add_table(rows=len(items), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER if center else WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    for idx, (label, value) in enumerate(items):
        left = table.cell(idx, 0)
        right = table.cell(idx, 1)
        left.text = label
        right.text = value
        left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(left, ACCENT_LIGHT)
        set_cell_border(left)
        set_cell_border(right)
        for cell in (left, right):
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.name = 'Arial' if cell is left else 'Georgia'
                    run.font.size = Pt(9.5)
                    if cell is left:
                        run.bold = True
                        run.font.color.rgb = ACCENT
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


def add_title_block(doc: Document, title: str, kind: str, metadata: list[tuple[str, str]], lead_text: str | None) -> None:
    title_par = doc.add_paragraph(style='Doc Title')
    title_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_par.add_run(title)

    if metadata:
        add_metadata_table(doc, metadata, center=(kind == 'report'))
    if lead_text:
        lead = doc.add_paragraph(style='Lead Paragraph')
        lead.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lead.add_run(lead_text)



def style_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for ridx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_border(cell)
            if ridx == 0:
                set_cell_shading(cell, ACCENT_LIGHT)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(3)
                for run in p.runs:
                    run.font.name = 'Arial' if ridx == 0 else 'Georgia'
                    run.font.size = Pt(9 if ridx == 0 else 9.5)
                    if ridx == 0:
                        run.bold = True
                        run.font.color.rgb = ACCENT


def render_table(doc: Document, tag: Tag) -> None:
    rows = tag.find_all('tr')
    if not rows:
        return
    col_count = max(len(r.find_all(['th', 'td'])) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    for ridx, row in enumerate(rows):
        cells = row.find_all(['th', 'td'])
        for cidx in range(col_count):
            cell = table.cell(ridx, cidx)
            cell.text = ''
            if cidx < len(cells):
                cell.text = cells[cidx].get_text(' ', strip=True)
    style_table(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def render_list(doc: Document, tag: Tag, ordered: bool, level: int = 0) -> None:
    for item in tag.find_all('li', recursive=False):
        style_name = ('List Number' if ordered else 'List Bullet') + (f' {level + 1}' if level else '')
        if style_name not in doc.styles:
            style_name = 'List Number' if ordered else 'List Bullet'
        para = doc.add_paragraph(style=style_name)
        para.paragraph_format.space_after = Pt(3)
        para.paragraph_format.left_indent = Inches(0.18 * level)
        nested: list[Tag] = []
        for child in item.children:
            if isinstance(child, Tag) and child.name in {'ul', 'ol'}:
                nested.append(child)
            elif isinstance(child, Tag) and child.name == 'p':
                for sub in child.children:
                    append_inline(para, sub)
            else:
                append_inline(para, child)
        for nested_list in nested:
            render_list(doc, nested_list, nested_list.name == 'ol', level + 1)


def render_blocks(doc: Document, nodes: list[Tag]) -> None:
    for node in nodes:
        if isinstance(node, NavigableString):
            if not str(node).strip():
                continue
            p = doc.add_paragraph()
            p.add_run(str(node).strip())
            continue
        if not isinstance(node, Tag):
            continue

        name = node.name.lower()
        if name == 'h2':
            doc.add_paragraph(node.get_text(' ', strip=True), style='Heading 1')
        elif name == 'h3':
            doc.add_paragraph(node.get_text(' ', strip=True), style='Heading 2')
        elif name == 'h4':
            doc.add_paragraph(node.get_text(' ', strip=True), style='Heading 3')
        elif name == 'p':
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(7)
            for child in node.children:
                append_inline(p, child)
        elif name == 'ul':
            render_list(doc, node, ordered=False)
        elif name == 'ol':
            render_list(doc, node, ordered=True)
        elif name == 'table':
            render_table(doc, node)
        elif name == 'blockquote':
            paragraphs = node.find_all('p', recursive=False) or [node]
            for quote in paragraphs:
                p = doc.add_paragraph(style='Quote Custom')
                for child in quote.children:
                    append_inline(p, child)
        elif name == 'pre':
            p = doc.add_paragraph(style='Code Block')
            p.add_run(node.get_text())
            set_paragraph_border(p, color=GRAY_BORDER)
        elif name == 'hr':
            doc.add_page_break()


def build_docx(md_path: Path) -> None:
    text = md_path.read_text()
    soup = markdown_to_soup(text)
    body = soup.body or soup
    nodes = [child for child in body.children if not (isinstance(child, NavigableString) and not str(child).strip())]
    if not nodes:
        return

    title = nodes[0].get_text(' ', strip=True) if isinstance(nodes[0], Tag) and nodes[0].name == 'h1' else md_path.stem
    kind = 'report' if 'Report' in md_path.name else 'one_pager' if 'One_Pager' in md_path.name else 'script'

    metadata: list[tuple[str, str]] = []
    lead_text: str | None = None
    start_index = 1 if isinstance(nodes[0], Tag) and nodes[0].name == 'h1' else 0

    if start_index < len(nodes) and isinstance(nodes[start_index], Tag) and nodes[start_index].name == 'p':
        parsed = parse_metadata_paragraph(nodes[start_index])
        if parsed:
            metadata = parsed
            start_index += 1
        elif kind == 'one_pager':
            lead_text = nodes[start_index].get_text(' ', strip=True)
            start_index += 1

    doc = Document()
    configure_document(doc, title, kind)
    add_title_block(doc, title, kind, metadata, lead_text)
    render_blocks(doc, nodes[start_index:])
    doc.save(md_path.with_suffix('.docx'))


def main() -> None:
    for name in FILES:
        build_docx(BASE / name)


if __name__ == '__main__':
    main()
